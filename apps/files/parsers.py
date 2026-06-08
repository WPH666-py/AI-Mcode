import os
import zipfile
import io
import csv


def parse_file(file_path, file_type):
    parsed = ""

    if file_type == 'docx':
        parsed = _parse_docx(file_path)
    elif file_type == 'pdf':
        parsed = _parse_pdf(file_path)
    elif file_type in ('xlsx', 'xls'):
        parsed = _parse_xlsx(file_path)
    elif file_type == 'csv':
        parsed = _parse_csv(file_path)
    elif file_type == 'txt':
        parsed = _parse_txt(file_path)
    elif file_type == 'md':
        parsed = _parse_txt(file_path)
    elif file_type in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'):
        parsed = _parse_image(file_path)
    elif file_type == 'zip':
        parsed = _parse_zip(file_path)
    else:
        parsed = f"[不支持的文件类型: {file_type}]"

    return parsed


def _parse_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(' | '.join(cells))
        return '\n'.join(lines)
    except ImportError:
        return "[python-docx 未安装]"


def _parse_pdf(file_path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        lines = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
        return '\n'.join(lines)
    except ImportError:
        return "[PyPDF2 未安装]"


def _parse_xlsx(file_path):
    try:
        import pandas as pd
        xls = pd.ExcelFile(file_path)
        lines = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            lines.append(f"--- Sheet: {sheet_name} ---")
            lines.append(df.to_csv(index=False))
        return '\n'.join(lines)
    except ImportError:
        return "[pandas/openpyxl 未安装]"


def _parse_csv(file_path):
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        return df.to_csv(index=False)
    except ImportError:
        return "[pandas 未安装]"


def _parse_txt(file_path):
    from . import _fast_read_text as _frt
    if _frt is not None:
        return _frt(file_path)
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _parse_image(file_path):
    return f"[图片文件: {os.path.basename(file_path)}]"


def _parse_zip(file_path):
    results = []
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            for name in zf.namelist():
                ext = name.split('.')[-1].lower() if '.' in name else ''
                try:
                    content = zf.read(name)
                    if ext in ('txt', 'md', 'csv'):
                        results.append(f"--- {name} ---")
                        results.append(content.decode('utf-8', errors='ignore'))
                    elif ext in ('docx', 'pdf', 'xlsx', 'xls'):
                        tmp_path = os.path.join(os.path.dirname(file_path), f"_tmp_{name}")
                        with open(tmp_path, 'wb') as f:
                            f.write(content)
                        results.append(f"--- {name} ---")
                        results.append(parse_file(tmp_path, ext))
                        os.remove(tmp_path)
                    else:
                        results.append(f"--- {name} (二进制文件) ---")
                except Exception as e:
                    results.append(f"--- {name} (解析失败: {e}) ---")
        return '\n'.join(results)
    except ImportError:
        return "[zipfile 模块错误]"
